import os
import tempfile
import PyPDF2
import streamlit as st
import torch
import time
import pandas as pd
import warnings
import psutil
import docx
import datetime
import bson

from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import OllamaLLM
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS, Chroma
from langchain.chains import LLMChain
from langchain.docstore.document import Document
from langchain.prompts import PromptTemplate
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain_community.tools import DuckDuckGoSearchRun
from utils import remove_directory_recursively, get_quantum_flag

warnings.filterwarnings("ignore", category=DeprecationWarning)
os.environ["PYTHONWARNINGS"] = "ignore::DeprecationWarning"

# -----------------------------
# Quantum modules (stubs/ready)
# -----------------------------

class QuantumEmbeddingEncoder:
    """
    Pluggable quantum embeddings.
    Default: uses classical embeddings, but the wrapper lets you swap in
    quantum feature maps/kernels later without touching call sites.
    """
    def __init__(self, base_embeddings: HuggingFaceEmbeddings):
        self.base = base_embeddings

    def embed_documents(self, docs):
        # TODO: replace with quantum kernel embeddings (feature maps / tensor networks)
        return self.base.embed_documents(docs)

    def embed_query(self, query: str):
        # TODO: replace with quantum kernel embedding for queries
        return self.base.embed_query(query)
    def __call__(self, text: str):
        """Allows the object to be called like a function, for embedding single queries."""
        return self.embed_query(text)

def quantum_optimize_context(chunks, query, token_budget=1500, prefer_diversity=True):
    """
    Quantum optimization stub for selecting the best chunk set under a token budget.
    Replace with QAOA/annealing backend. Maintains same input/output API for drop-in replacement.
    """
    selected = []
    total_tokens = 0
    seen_sources = set()

    def score(c):
        # Placeholder objective; real version uses kernel overlaps or learned objective
        return -len(c.page_content)

    for c in sorted(chunks, key=score):
        approx_tokens = max(1, len(c.page_content) // 4)
        if total_tokens + approx_tokens > token_budget:
            continue
        if prefer_diversity:
            src = c.metadata.get("source", "unknown")
            if src in seen_sources and len(selected) >= 2:
                continue
            seen_sources.add(src)
        selected.append(c)
        total_tokens += approx_tokens
        if total_tokens >= token_budget:
            break
    return selected

def quantum_rerank_answers(initial_answer: str, query: str, candidates=None):
    """
    Quantum-assisted rerank stub. In production, evaluate candidate answers with
    variational circuits / quantum Boltzmann machines and pick the highest-scoring.
    """
    return initial_answer

# -----------------------------
# EnhancedRAG (Quantum-ready)
# -----------------------------

class EnhancedRAG:
    def __init__(self, llm_model_name="llama3.2:latest",
                 embedding_model_name="sentence-transformers/all-MiniLM-L6-v2",
                 chunk_size=1000, chunk_overlap=200, use_gpu=True):
        """Initialize the Enhanced RAG system with multiple modes and quantum-ready hooks."""
        self.llm_model_name = llm_model_name
        self.embedding_model_name = embedding_model_name
        self.use_gpu = use_gpu and torch.cuda.is_available()
        self.temp_dirs = []
        self.device = "cuda" if self.use_gpu else "cpu"
        st.sidebar.info(f"Using device: {self.device}")

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size, chunk_overlap=chunk_overlap, length_function=len
        )

        # Embeddings: base + optional quantum wrapper
        try:
            base_embeddings = HuggingFaceEmbeddings(
                model_name=embedding_model_name,
                model_kwargs={"device": self.device}
            )
            if get_quantum_flag("enable_quantum_embeddings"):
                self.embeddings = QuantumEmbeddingEncoder(base_embeddings)
                st.sidebar.success(f"Quantum embeddings enabled on {embedding_model_name}")
            else:
                self.embeddings = base_embeddings
                st.sidebar.success(f"Embeddings model loaded: {embedding_model_name}")
        except Exception as e:
            st.sidebar.error(f"Failed to load embeddings model: {str(e)}")
            self.embeddings = None

        # LLM
        try:
            callbacks = [StreamingStdOutCallbackHandler()]
            self.llm = OllamaLLM(model=llm_model_name, callbacks=callbacks)
            st.sidebar.success(f"LLM loaded: {llm_model_name}")
        except Exception as e:
            st.sidebar.error(f"Failed to load LLM: {str(e)}")
            self.llm = None

        self.doc_vector_store = None
        self.web_vector_store = None
        self.documents_processed = 0
        self.processing_times = {}
        self.sources = []
        self.errors = []

    def __del__(self):
        """Cleanup temporary directories when object is garbage collected."""
        for temp_dir in self.temp_dirs:
            try:
                if os.path.exists(temp_dir):
                    remove_directory_recursively(temp_dir)
            except:
                pass

    # -----------------------------
    # 1) Document Ingestion & Embedding (quantum-ready)
    # -----------------------------
    def process_files(self, files, user_id=None, mongodb=None, notebook_id=None, is_nested=False, domains=None):
        """Process files and build vector store."""
        if self.embeddings is None:
            st.error("Embeddings model not initialized. Unable to process files.")
            return False

        all_docs = []
        document_metadata = []
        status_msg = st.empty()
        status_msg.info("Processing files...")

        if "temp_dir" not in st.session_state:
            st.session_state["temp_dir"] = None
        temp_dir = tempfile.mkdtemp()
        self.temp_dirs.append(temp_dir)
        st.session_state["temp_dir"] = temp_dir

        start_time = time.time()
        mem_before = psutil.virtual_memory().used / (1024 * 1024 * 1024)
        total_files = len(files)

        for i, file in enumerate(files):
            try:
                status_msg.info(f"Processing {file.name} ({i+1}/{total_files})...")
                file_start_time = time.time()

                file_type = "unknown"
                if file.name.lower().endswith('.pdf'):
                    file_type = "pdf"
                elif file.name.lower().endswith(('.docx', '.doc')):
                    file_type = "docx"
                elif file.name.lower().endswith('.txt'):
                    file_type = "txt"

                file_path = os.path.join(temp_dir, file.name)
                file.seek(0)
                file_content = file.read()
                with open(file_path, "wb") as f:
                    f.write(file_content)

                text = ""
                page_count = 0

                if file_type == "pdf":
                    try:
                        with open(file_path, "rb") as f:
                            pdf = PyPDF2.PdfReader(f)
                            page_count = len(pdf.pages)
                            for page in pdf.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    text += page_text + "\n\n"
                    except Exception as e:
                        st.error(f"Error extracting text from PDF {file.name}: {str(e)}")
                        continue

                elif file_type == "docx":
                    try:
                        doc = docx.Document(file_path)
                        page_count = len(doc.paragraphs)
                        for para in doc.paragraphs:
                            text += para.text + "\n\n"
                    except Exception as e:
                        st.error(f"Error extracting text from DOCX {file.name}: {str(e)}")
                        continue

                elif file_type == "txt":
                    try:
                        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                            text = f.read()
                            page_count = text.count("\n") + 1
                    except Exception as e:
                        st.error(f"Error extracting text from TXT {file.name}: {str(e)}")
                        continue

                if not text.strip():
                    st.warning(f"No text content found in {file.name}. Skipping...")
                    continue

                docs = [Document(page_content=text, metadata={
                    "source": file.name, "notebook_id": notebook_id, "file_type": file_type
                })]

                try:
                    split_docs = self.text_splitter.split_documents(docs)
                except Exception as e:
                    st.error(f"Error splitting document {file.name}: {str(e)}")
                    continue

                all_docs.extend(split_docs)
                processing_time = time.time() - file_start_time
                document_metadata.append({
                    "filename": file.name,
                    "file_type": file_type,
                    "page_count": page_count,
                    "chunk_count": len(split_docs),
                    "processing_time": processing_time,
                    "notebook_id": notebook_id
                })
                st.sidebar.success(f"Processed {file.name}: {len(split_docs)} chunks in {processing_time:.2f}s")
                self.processing_times[file.name] = {"chunks": len(split_docs), "time": processing_time}

            except Exception as e:
                error_msg = f"Error processing {file.name}: {str(e)}"
                self.errors.append(error_msg)
                st.sidebar.error(error_msg)

        if all_docs:
            status_msg.info("Building vector index...")
            try:
                index_start_time = time.time()
                # FAISS backend; embeddings path can be quantum-enabled via QuantumEmbeddingEncoder
                self.doc_vector_store = FAISS.from_documents(all_docs, self.embeddings)
                index_end_time = time.time()

                mem_after = psutil.virtual_memory().used / (1024 * 1024 * 1024)
                mem_used = mem_after - mem_before
                total_time = time.time() - start_time

                status_msg.success(f"Completed processing {len(all_docs)} chunks in {total_time:.2f}s")
                self.processing_times["index_building"] = index_end_time - index_start_time
                self.processing_times["total_time"] = total_time
                self.processing_times["memory_used_gb"] = mem_used
                self.documents_processed = len(all_docs)
                return True
            except Exception as e:
                error_msg = f"Error creating vector store: {str(e)}"
                self.errors.append(error_msg)
                st.error(error_msg)
                status_msg.error(error_msg)
                return False
        else:
            status_msg.error("No content extracted from files")
            return False

    # -----------------------------
    # Enhancement helper (unchanged logic, clearer prompt)
    # -----------------------------
    def enhance_answer(self, initial_answer, query, source_content):
        enhance_template = """
You are an expert content enhancer. Improve the answer while keeping facts grounded in the sources.

QUERY:
{query}

INITIAL ANSWER:
{initial_answer}

SOURCE CONTENT (EXTRACT):
{source_content}

Enhance for clarity, coverage, and structure. Maintain factual consistency.
ENHANCED ANSWER:
"""
        enhancement_prompt = PromptTemplate(
            template=enhance_template,
            input_variables=["query", "initial_answer", "source_content"]
        )
        enhancement_chain = LLMChain(llm=self.llm, prompt=enhancement_prompt)
        summarized_sources = "\n\n".join([
            (f"SOURCE {i+1}:\n{src[:500]}..." if len(src) > 500 else f"SOURCE {i+1}:\n{src}")
            for i, src in enumerate(source_content[:3])
        ])
        try:
            enhanced_result = enhancement_chain.invoke({
                "query": query, "initial_answer": initial_answer, "source_content": summarized_sources
            })
            return enhanced_result["text"].strip()
        except Exception as e:
            st.warning(f"Enhancement step issue: {str(e)}. Using initial answer.")
            self.errors.append(f"Enhancement error: {str(e)}")
            return initial_answer

    # -----------------------------
    # 2) Context Selection (quantum-ready) + 4) Re-ranking (quantum-ready)
    # -----------------------------
    def direct_retrieval_answer(self, query, user_id=None, mongodb=None, notebook_id=None):
        """Generate an answer using direct document retrieval."""
        if not hasattr(self, 'doc_vector_store') or not self.doc_vector_store:
            return "Please upload and process documents first."

        # Retrieve more than k to allow optimization to choose subset under budget
        docs = self.doc_vector_store.similarity_search(query, k=12)

        # Quantum context selection (pre-LLM)
        if get_quantum_flag("enable_quantum_context_selection"):
            selected_docs = quantum_optimize_context(docs, query, token_budget=1800, prefer_diversity=True)
        else:
            selected_docs = docs[:4]

        if not selected_docs:
            return "No relevant information found in your documents."

        source_content = []
        for doc in selected_docs:
            source_content.append({
                "content": doc.page_content,
                "source": doc.metadata.get("source", "Unknown"),
                "file_type": doc.metadata.get("file_type", "document")
            })

        context_text = " ".join([doc.page_content for doc in selected_docs])
        prompt = f"""
Answer the following question based on the provided context from documents.

Question: {query}

Context:
{context_text}

Provide a clear, concise answer strictly grounded in the context.
"""
        initial_answer = self.llm(prompt)

        # Quantum-assisted rerank (post-LLM)
        final_answer = quantum_rerank_answers(initial_answer, query) if get_quantum_flag("enable_quantum_rerank") else initial_answer

        return {"answer": final_answer, "sources": source_content}

    def enhanced_rag_answer(self, query, user_id=None, mongodb=None, notebook_id=None):
        """Generate an answer using the enhanced RAG pipeline with optional quantum steps."""
        if hasattr(self, 'doc_vector_store') and self.doc_vector_store:
            docs = self.doc_vector_store.similarity_search(query, k=12)
            if get_quantum_flag("enable_quantum_context_selection"):
                docs = quantum_optimize_context(docs, query, token_budget=2000, prefer_diversity=True)
            else:
                docs = docs[:4]

            source_content = []
            for doc in docs:
                source_content.append({
                    "content": doc.page_content,
                    "source": doc.metadata.get("source", "Unknown"),
                    "file_type": doc.metadata.get("file_type", "document")
                })

            prompt = f"""Based on the following context, answer the question:
{query}

Context:
{' '.join([d.page_content for d in docs])}

Answer:"""
            initial_answer = self.llm(prompt)
            enhanced_answer = self.enhance_answer(initial_answer, query, [d["content"] for d in source_content])

            if get_quantum_flag("enable_quantum_rerank"):
                enhanced_answer = quantum_rerank_answers(enhanced_answer, query)

            return {
                "answer": enhanced_answer,
                "initial_answer": initial_answer,
                "sources": source_content
            }
        else:
            return {"answer": "Please upload and process documents first."}

    # -----------------------------
    # 3) Hybrid mode (doc side quantum context optional)
    # -----------------------------
    def hybrid_answer(self, query, user_id=None, mongodb=None, notebook_id=None):
        """Generate an answer using hybrid search combining documents and web search."""
        doc_sources = []
        web_sources = []
        combined_context = ""
        
        has_documents = hasattr(self, 'doc_vector_store') and self.doc_vector_store
        
        if has_documents:
            docs = self.doc_vector_store.similarity_search(query, k=3)
            
            for doc in docs:
                doc_content = doc.page_content
                doc_source = doc.metadata.get("source", "Unknown document")
                
                combined_context += f"Document: {doc_source}\nContent: {doc_content}\n\n"
                
                doc_sources.append({
                    "content": doc_content,
                    "source": doc_source,
                    "file_type": doc.metadata.get("file_type", "document")
                })
        
        web_results = self.simulate_search(query, num_results=3)
        
        for result in web_results:
            combined_context += f"Web: {result['title']}\nContent: {result['content']}\n\n"
            
            web_sources.append({
                "content": result['content'],
                "source": result['title'],
                "file_type": "web"
            })
        
        all_sources = doc_sources + web_sources
        
        prompt = f"""
        I need to answer the following question thoroughly: {query}
        
        I have collected the following information:
        
        {combined_context}
        
        Based on all this information, provide a comprehensive, well-structured answer. 
        Integrate information from both documents and web sources when available.
        """
        
        answer = self.llm(prompt)
        
        return {
            "answer":answer,
            "sources": all_sources,
            "doc_sources_count": len(doc_sources),
            "web_sources_count": len(web_sources)
        }

    # ADD THIS NEW METHOD TO SUPPORT hybrid_answer
    def simulate_search(self, query: str, num_results: int = 3):
        """A placeholder for a web search function."""
        # In a real application, this would use an API like Google Search, Bing, etc.
        # For now, it returns dummy data to prevent errors.
        return [
            {
                "title": f"Simulated Web Result 1 for '{query}'",
                "content": "This is the content from the first simulated web search result. It provides context related to the user's query."
            },
            {
                "title": f"Simulated Web Result 2 for '{query}'",
                "content": "This is another piece of information found on the web, designed to complement the document sources."
            }
        ]

    # YOUR ASK METHOD (WITH THE UNNECESSARY UI CODE REMOVED)
    def ask(self, query, mode="direct_retrieval", user_id=None, mongodb=None, notebook_id=None):
        """Ask a question and get an answer from the RAG system."""
        if not hasattr(self, 'doc_vector_store') or not self.doc_vector_store:
            return "Please upload and process documents first."
        
        start_time = time.time()
        
        if mongodb and user_id:
            mongodb.log_query(user_id, query, 0, notebook_id)
        
        try:
            if mode == "enhanced_rag":
                result = self.enhanced_rag_answer(query, user_id, mongodb, notebook_id)
            elif mode == "hybrid":
                result = self.hybrid_answer(query, user_id, mongodb, notebook_id)
            else:
                result = self.direct_retrieval_answer(query, user_id, mongodb, notebook_id)
                
            query_time = time.time() - start_time
            
            if isinstance(result, dict):
                result["query_time"] = query_time
                result["mode"] = mode
            
            if mongodb and user_id:
                mongodb.log_query(user_id, query, query_time, notebook_id)
                
            return result
            
        except Exception as e:
            error_message = f"Error processing query: {str(e)}"
            print(error_message) # Good for debugging on the server
            return error_message